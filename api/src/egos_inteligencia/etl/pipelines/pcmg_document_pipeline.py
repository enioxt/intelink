"""
PCMG Document Pipeline — EGOS Inteligência
Processa documentos policiais: PDF, DOCX, DOC, XLSX, XLS, CSV, TXT

Para: comunicações, ordens de serviço, ocorrências, inquéritos,
      relatórios, perícias, planilhas operacionais

Sacred Code: 000.111.369.963.1618
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
import os
import re
import subprocess
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

import aiofiles

logger = logging.getLogger(__name__)


@dataclass
class DocumentProcessingResult:
    """Resultado do processamento de documento."""
    source_path: str
    file_type: str  # pdf, docx, xlsx, csv, txt
    raw_text: str | None
    structured_data: dict[str, Any] | None
    tables: list[list[list[str]]]  # Tabelas extraídas
    analysis: dict[str, Any] | None
    metadata: dict[str, Any]
    success: bool
    errors: list[str] = field(default_factory=list)


@dataclass
class PoliceDocument:
    """Documento policial estruturado."""
    document_type: str  # ocorrencia, inquerito, ordem_servico, etc.
    document_number: str | None  # Número REDS, inquérito, etc.
    date: str | None
    location: str | None
    author: str | None  # Quem registrou
    involved_parties: list[dict[str, Any]]  # Pessoas com papéis
    facts: str | None  # Narração dos fatos
    legal_basis: list[str]  # Fundamentação legal
    recommendations: list[str]  # Providências
    attachments: list[str]  # Anexos mencionados
    raw_extracted: str  # Texto bruto extraído


class PCMGDocumentPipeline:
    """
    Pipeline para documentos policiais em múltiplos formatos.
    
    Fluxo:
    1. Detectar tipo de arquivo
    2. Extrair texto/tabelas (ferramenta específica por formato)
    3. Estruturar dados (detectar campos policiais)
    4. Analisar contexto
    5. Validar e salvar
    
    Suporta: PDF, DOCX, DOC, XLSX, XLS, CSV, TXT
    """
    
    SUPPORTED_FORMATS = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".xlsx": "xlsx",
        ".xls": "xls",
        ".csv": "csv",
        ".txt": "txt",
        ".md": "txt",
        ".rtf": "txt",
    }
    
    # Padrões para campos de documentos policiais
    FIELD_PATTERNS = {
        "reds": [
            r"(?:REDS|Reds|redes)[\s:\-]*(\d{7}/\d{4})",
            r"(?:REDS|Reds)[\s:\-]*(\d{7})/(\d{4})",
        ],
        "inquerito": [
            r"(?:Inquérito|IP|inquerito)[\s:\-]*(\d{3,5}/\d{4})",
            r"IP[\s:\-]*(\d{3,5})/(\d{4})",
        ],
        "data": [
            r"(?:Data|em)\s*[:\-]?\s*(\d{2}/\d{2}/\d{4})",
            r"(?:Belo Horizonte|BH|MG)[,\s]+(\d{2})\s+de\s+([a-zç]+)\s+de\s+(\d{4})",
        ],
        "delegacia": [
            r"(?:Delegacia|DP|DPCAMI)\s+de\s+([A-Za-z\s]+)",
            r"(?:Delegacia|DP)\s+([A-Za-z\s\d]+)",
        ],
        "escrivao": [
            r"(?:Escrivão|Esc\.|Esc)[\s:\-]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        ],
        "delegado": [
            r"(?:Delegado|Dr\.|Dra\.)[\s:\-]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        ],
        "vitima": [
            r"(?:VÍTIMA|Vítima|vitima|LESADO)[\s:\-]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        ],
        "indiciado": [
            r"(?:INDICIADO|Indiciado|indiciado|INVESTIGADO)[\s:\-]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        ],
        "local": [
            r"(?:Local do Fato|Local)[\s:\-]+([A-Z][a-zA-Z0-9\s,\.]+?)(?:\n|Data|$)",
        ],
    }
    
    def __init__(self, temp_dir: str | None = None):
        self.temp_dir = temp_dir or "/tmp/pcmg_docs"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    async def process_document(
        self,
        file_path: str,
        extract_tables: bool = True,
        analyze: bool = True,
        progress_callback: Callable[[str, float], None] | None = None,
    ) -> DocumentProcessingResult:
        """
        Processa um documento policial completo.
        
        Args:
            file_path: Caminho do arquivo
            extract_tables: Se deve extrair tabelas (Excel/CSV)
            analyze: Se deve analisar contexto policial
            progress_callback: Callback (etapa, progresso)
        
        Returns:
            DocumentProcessingResult com todos os dados extraídos
        """
        result = DocumentProcessingResult(
            source_path=file_path,
            file_type=self._detect_file_type(file_path),
            raw_text=None,
            structured_data=None,
            tables=[],
            analysis=None,
            metadata={},
            success=False,
            errors=[],
        )
        
        try:
            # Fase 1: Extrair texto
            if progress_callback:
                await asyncio.get_event_loop().run_in_executor(
                    None, lambda: progress_callback("extract_text", 0.0)
                )
            
            raw_text = await self._extract_text(file_path, result.file_type)
            result.raw_text = raw_text
            result.metadata["char_count"] = len(raw_text) if raw_text else 0
            
            if progress_callback:
                await asyncio.get_event_loop().run_in_executor(
                    None, lambda: progress_callback("extract_text", 1.0)
                )
            
            # Fase 2: Extrair tabelas (se aplicável)
            if extract_tables and result.file_type in ["xlsx", "xls", "csv"]:
                if progress_callback:
                    await asyncio.get_event_loop().run_in_executor(
                        None, lambda: progress_callback("extract_tables", 0.0)
                    )
                
                tables = await self._extract_tables(file_path, result.file_type)
                result.tables = tables
                
                if progress_callback:
                    await asyncio.get_event_loop().run_in_executor(
                        None, lambda: progress_callback("extract_tables", 1.0)
                    )
            
            # Fase 3: Estruturar dados policiais
            if progress_callback:
                await asyncio.get_event_loop().run_in_executor(
                    None, lambda: progress_callback("structure", 0.0)
                )
            
            structured = self._structure_police_data(raw_text, result.tables)
            result.structured_data = structured
            
            if progress_callback:
                await asyncio.get_event_loop().run_in_executor(
                    None, lambda: progress_callback("structure", 1.0)
                )
            
            # Fase 4: Analisar contexto
            if analyze and raw_text:
                if progress_callback:
                    await asyncio.get_event_loop().run_in_executor(
                        None, lambda: progress_callback("analyze", 0.0)
                    )
                
                analysis = self._analyze_document_context(raw_text, structured)
                result.analysis = analysis
                
                if progress_callback:
                    await asyncio.get_event_loop().run_in_executor(
                        None, lambda: progress_callback("analyze", 1.0)
                    )
            
            result.success = True
            
        except Exception as e:
            logger.exception(f"Erro processando documento {file_path}")
            result.errors.append(str(e))
        
        return result
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detecta tipo de arquivo pela extensão."""
        ext = Path(file_path).suffix.lower()
        return self.SUPPORTED_FORMATS.get(ext, "unknown")
    
    async def _extract_text(self, file_path: str, file_type: str) -> str | None:
        """Extrai texto bruto do arquivo usando ferramenta apropriada."""
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "doc": self._extract_doc,
            "xlsx": self._extract_spreadsheet_text,
            "xls": self._extract_spreadsheet_text,
            "csv": self._extract_csv_text,
            "txt": self._extract_txt,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Tipo de arquivo não suportado: {file_type}")
        
        return await extractor(file_path)
    
    async def _extract_pdf(self, pdf_path: str) -> str:
        """Extrai texto de PDF usando pdftotext ou pymupdf."""
        # Tentar pdftotext primeiro (mais rápido)
        try:
            cmd = ["pdftotext", "-layout", pdf_path, "-"]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                check=True
            )
            if result.stdout.strip():
                return result.stdout
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        
        # Fallback para PyMuPDF
        try:
            import fitz  # PyMuPDF
            
            text_parts = []
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text_parts.append(page.get_text())
            
            return "\n".join(text_parts)
        except ImportError:
            raise RuntimeError(
                "PDF extraction failed. Install: sudo apt install poppler-utils "
                "ou pip install PyMuPDF"
            )
    
    async def _extract_docx(self, docx_path: str) -> str:
        """Extrai texto de DOCX usando python-docx."""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Também extrair de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))
            
            return "\n".join(paragraphs)
            
        except ImportError:
            raise RuntimeError("python-docx não instalado: pip install python-docx")
    
    async def _extract_doc(self, doc_path: str) -> str:
        """Extrai texto de DOC usando antiword ou conversão."""
        # Tentar antiword
        try:
            cmd = ["antiword", doc_path]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                check=True
            )
            return result.stdout
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        
        # Fallback: converter para docx primeiro
        try:
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", self.temp_dir,
                doc_path
            ]
            subprocess.run(cmd, capture_output=True, timeout=120, check=True)
            
            # Encontrar arquivo convertido
            base_name = Path(doc_path).stem
            converted = Path(self.temp_dir) / f"{base_name}.docx"
            
            if converted.exists():
                text = await self._extract_docx(str(converted))
                converted.unlink()  # Limpar
                return text
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        
        raise RuntimeError(
            "DOC extraction failed. Install: sudo apt install antiword "
            "ou libreoffice-writer"
        )
    
    async def _extract_spreadsheet_text(self, xlsx_path: str) -> str:
        """Extrai texto de planilhas como texto formatado."""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            text_parts = []
            
            for sheet in wb.worksheets:
                text_parts.append(f"=== {sheet.title} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):  # Pular linhas vazias
                        text_parts.append(" | ".join(row_text))
                
                text_parts.append("")
            
            return "\n".join(text_parts)
            
        except ImportError:
            raise RuntimeError("openpyxl não instalado: pip install openpyxl")
    
    async def _extract_csv_text(self, csv_path: str) -> str:
        """Extrai texto de CSV."""
        text_parts = []
        
        async with aiofiles.open(csv_path, "r", encoding="utf-8", errors="ignore") as f:
            content = await f.read()
        
        reader = csv.reader(io.StringIO(content))
        for row in reader:
            text_parts.append(" | ".join(row))
        
        return "\n".join(text_parts)
    
    async def _extract_txt(self, txt_path: str) -> str:
        """Lê arquivo de texto."""
        async with aiofiles.open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            return await f.read()
    
    async def _extract_tables(
        self,
        file_path: str,
        file_type: str,
    ) -> list[list[list[str]]]:
        """Extrai tabelas de planilhas."""
        if file_type in ["xlsx", "xls"]:
            return await self._extract_excel_tables(file_path)
        elif file_type == "csv":
            return await self._extract_csv_tables(file_path)
        
        return []
    
    async def _extract_excel_tables(self, xlsx_path: str) -> list[list[list[str]]]:
        """Extrai tabelas de Excel."""
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            tables = []
            
            for sheet in wb.worksheets:
                table = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):
                        table.append(row_data)
                
                if table:
                    tables.append(table)
            
            return tables
            
        except ImportError:
            return []
    
    async def _extract_csv_tables(self, csv_path: str) -> list[list[list[str]]]:
        """Extrai tabela de CSV."""
        table = []
        
        async with aiofiles.open(csv_path, "r", encoding="utf-8", errors="ignore") as f:
            content = await f.read()
        
        reader = csv.reader(io.StringIO(content))
        for row in reader:
            table.append(row)
        
        return [table] if table else []
    
    def _structure_police_data(
        self,
        raw_text: str | None,
        tables: list[list[list[str]]],
    ) -> PoliceDocument | None:
        """
        Estrutura dados brutos em documento policial.
        
        Detecta campos específicos de documentos policiais.
        """
        if not raw_text:
            return None
        
        text = raw_text
        
        # Extrair campos usando padrões
        doc_number = self._extract_field(text, "reds") or self._extract_field(text, "inquerito")
        date = self._extract_field(text, "data")
        location = self._extract_field(text, "local") or self._extract_field(text, "delegacia")
        author = self._extract_field(text, "escrivao") or self._extract_field(text, "delegado")
        
        # Detectar tipo de documento
        doc_type = self._detect_document_type(text)
        
        # Extrair partes envolvidas
        involved = self._extract_involved_parties(text)
        
        # Extrair fatos (narração)
        facts = self._extract_facts(text, doc_type)
        
        # Extrair fundamentação legal
        legal = self._extract_legal_references(text)
        
        # Extrair providências/recomendações
        recommendations = self._extract_recommendations(text)
        
        # Extrair anexos mencionados
        attachments = self._extract_attachments(text)
        
        return PoliceDocument(
            document_type=doc_type,
            document_number=doc_number,
            date=date,
            location=location,
            author=author,
            involved_parties=involved,
            facts=facts,
            legal_basis=legal,
            recommendations=recommendations,
            attachments=attachments,
            raw_extracted=text,
        )
    
    def _extract_field(self, text: str, field_type: str) -> str | None:
        """Extrai campo específico usando padrões."""
        patterns = self.FIELD_PATTERNS.get(field_type, [])
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _detect_document_type(self, text: str) -> str:
        """Detecta tipo de documento policial."""
        text_lower = text.lower()
        
        # Prioridades
        if "boletim de ocorrência" in text_lower or "ocorrência policial" in text_lower:
            return "ocorrencia_policial"
        
        if "inquérito policial" in text_lower or "termo de declaração" in text_lower:
            return "inquerito_policial"
        
        if "ordem de serviço" in text_lower or "ordem de missão" in text_lower:
            return "ordem_servico"
        
        if "comunicação de serviço" in text_lower or "passagem de serviço" in text_lower:
            return "comunicacao_servico"
        
        if "laudo pericial" in text_lower or "exame técnico" in text_lower:
            return "pericia_tecnica"
        
        if "relatório" in text_lower and "investigativo" in text_lower:
            return "relatorio_investigativo"
        
        return "documento_policia_civil"
    
    def _extract_involved_parties(self, text: str) -> list[dict[str, Any]]:
        """Extrai pessoas envolvidas e seus papéis."""
        parties = []
        
        roles = {
            "vitima": ["VÍTIMA", "Vítima", "vítima", "LESADO", "Lesado", "ofendido"],
            "indiciado": ["INDICIADO", "Indiciado", "indiciado", "INVESTIGADO", "autor"],
            "testemunha": ["TESTEMUNHA", "Testemunha", "testemunha"],
            "policial": ["POLICIAL", "Policial", "policial", "ESCRIVÃO", "DELEGADO"],
        }
        
        for role, keywords in roles.items():
            for keyword in keywords:
                # Padrão: KEYWORD: Nome ou KEYWORD - Nome
                pattern = rf"{keyword}[\s:\-\n]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)"
                for match in re.finditer(pattern, text):
                    name = match.group(1).strip()
                    if len(name) > 3 and name not in [p["name"] for p in parties]:
                        parties.append({
                            "name": name,
                            "role": role,
                            "source": "pattern",
                        })
        
        return parties
    
    def _extract_facts(self, text: str, doc_type: str) -> str | None:
        """Extrai narração dos fatos."""
        # Procurar por seção "DOS FATOS" ou similar
        patterns = [
            r"(?:DOS FATOS|DOS AUTOS|NARRATIVA|RELATO)[\s:]*\n+(.+?)(?:\n+(?:DAS?|DO)\s+[A-Z]|$)",
            r"(?:FATO|OCORRÊNCIA)[\s:]*\n+(.+?)(?:\n+|$)",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                facts = match.group(1).strip()
                if len(facts) > 50:
                    return facts[:2000]  # Limitar tamanho
        
        # Fallback: primeiras frases longas
        sentences = re.split(r"[.!?]+", text)
        for sent in sentences[:10]:
            if len(sent.strip()) > 100:
                return sent.strip()[:2000]
        
        return None
    
    def _extract_legal_references(self, text: str) -> list[str]:
        """Extrai referências legais."""
        refs = []
        
        # Artigos CP, CPP, Leis
        patterns = [
            r"(?:art\.?|artigo)\s+(\d+[\.,]?\d*)\s*,?\s*§?\s*(\d*)\s*(?:do)?\s*(Código Penal|CP|código de processo penal|CPP|CPP)",
            r"(?:Lei|lei)\s+(?:n?º?\s*)?(\d{4,5}[\./]\d{2,4})",
            r"(?:Constituição Federal|CF|CF/88)",
            r"(?:Estatuto|estatuto)\s+(?:do)?\s*([a-zA-Z\s]+)",
        ]
        
        for pattern in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                refs.append(match.group(0))
        
        return list(set(refs))
    
    def _extract_recommendations(self, text: str) -> list[str]:
        """Extrai providências recomendadas."""
        recs = []
        
        # Padrões de providências
        patterns = [
            r"(?:PROVIDÊNCIAS?|Providências?|CUMPRA-SE)[\s:]*\n+(.+?)(?:\n+|$)",
            r"(?:conclui-se|CONCLUI-SE)[\s:]*(.+?)(?:\.|$)",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                rec_text = match.group(1).strip()
                recs.extend([r.strip() for r in rec_text.split("\n") if r.strip()])
        
        return recs[:10]  # Top 10
    
    def _extract_attachments(self, text: str) -> list[str]:
        """Extrai anexos mencionados."""
        attachments = []
        
        pattern = r"(?:ANEXO|Anexo)[\s:]*(\w+)[\s:]*(?:[–\-])?\s*([^\n]+)"
        for match in re.finditer(pattern, text):
            attachments.append(f"{match.group(1)}: {match.group(2).strip()}")
        
        return attachments
    
    def _analyze_document_context(
        self,
        raw_text: str,
        structured: PoliceDocument | None,
    ) -> dict[str, Any]:
        """Analisa contexto do documento."""
        if not structured:
            return {"error": "Não foi possível estruturar documento"}
        
        # Detectar riscos
        risks = self._detect_risks(raw_text)
        
        # Calcular completude
        completeness = self._calculate_completeness(structured)
        
        # Sugerir ações
        actions = self._suggest_actions(structured, risks)
        
        # Verificar prazos
        deadlines = self._check_deadlines(structured)
        
        return {
            "risks": risks,
            "completeness_score": completeness,
            "completeness_details": {
                "has_number": structured.document_number is not None,
                "has_date": structured.date is not None,
                "has_location": structured.location is not None,
                "has_involved": len(structured.involved_parties) > 0,
                "has_facts": structured.facts is not None,
                "has_legal_basis": len(structured.legal_basis) > 0,
            },
            "suggested_actions": actions,
            "deadlines": deadlines,
            "classification": structured.document_type,
        }
    
    def _detect_risks(self, text: str) -> list[str]:
        """Detecta indicadores de risco."""
        risks = []
        
        risk_keywords = {
            "arma_de_fogo": ["arma", "revólver", "pistola", "fuzil", "escopeta"],
            "trafico_drogas": ["droga", "entorpecente", "cocaína", "maconha", "tráfico"],
            "violencia": ["agressão", "violência", "ameaça", "homicídio", "lesão"],
            "fuga": ["fuga", "evasão", "fugiu", "evadiu"],
            "organizada": ["organização criminosa", "facção", "milícia", "tráfico armado"],
        }
        
        text_lower = text.lower()
        for risk_type, keywords in risk_keywords.items():
            if any(kw in text_lower for kw in keywords):
                risks.append(risk_type)
        
        return risks
    
    def _calculate_completeness(self, doc: PoliceDocument) -> float:
        """Calcula score de completude do documento."""
        fields = [
            doc.document_number is not None,
            doc.date is not None,
            doc.location is not None,
            doc.author is not None,
            len(doc.involved_parties) > 0,
            doc.facts is not None,
            len(doc.legal_basis) > 0,
        ]
        
        return sum(fields) / len(fields)
    
    def _suggest_actions(self, doc: PoliceDocument, risks: list[str]) -> list[str]:
        """Sugere ações baseado no documento."""
        actions = []
        
        if doc.document_type == "ocorrencia_policial":
            actions.append("Registrar REDS no sistema")
            if "indiciado" in [p["role"] for p in doc.involved_parties]:
                actions.append("Consultar antecedentes criminais")
        
        if doc.document_type == "inquerito_policial":
            actions.append("Verificar prazo de conclusão")
            actions.append("Atualizar status de diligências")
        
        if "arma_de_fogo" in risks:
            actions.append("⚠️ Protocolo de apreensão de arma")
        
        if "trafico_drogas" in risks:
            actions.append("⚠️ Protocolo de apreensão de drogas")
        
        if doc.document_type == "ordem_servico":
            actions.append("Confirmar cumprimento da missão")
        
        return actions
    
    def _check_deadlines(self, doc: PoliceDocument) -> dict[str, Any]:
        """Verifica prazos baseado no tipo de documento."""
        deadlines = {}
        
        if doc.document_type == "inquerito_policial":
            deadlines["note"] = "Inquérito: prazo padrão 30 dias (prorrogável)"
            deadlines["urgency"] = "medium"
        
        if doc.document_type == "ordem_servico":
            deadlines["note"] = "Ordem de serviço: verificar data de cumprimento"
            deadlines["urgency"] = "high" if "urgente" in (doc.raw_extracted or "").lower() else "normal"
        
        return deadlines


# Instância global
document_pipeline = PCMGDocumentPipeline()


async def process_single_document(file_path: str) -> DocumentProcessingResult:
    """Processa um único documento."""
    return await document_pipeline.process_document(file_path)


async def batch_process_documents(
    file_paths: list[str],
    progress_callback: Callable[[int, int, str], None] | None = None,
) -> list[DocumentProcessingResult]:
    """Processa múltiplos documentos."""
    results = []
    total = len(file_paths)
    
    for i, path in enumerate(file_paths):
        if progress_callback:
            progress_callback(i + 1, total, f"Processando {Path(path).name}...")
        
        result = await process_single_document(path)
        results.append(result)
    
    return results
